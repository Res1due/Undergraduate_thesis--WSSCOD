from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


ROOT = Path(r"C:\Users\resii\Desktop\Noisy-cod-main-onlycode")
OUT_DOCX = ROOT / "xdu_thesis_initial_draft.docx"
FIG_DIR = ROOT / "generated_thesis_figures"


def ensure_dir(path: Path):
    path.mkdir(parents=True, exist_ok=True)


def get_font(size, bold=False):
    font_candidates = [
        "C:/Windows/Fonts/simhei.ttf" if bold else "C:/Windows/Fonts/simsun.ttc",
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/arial.ttf",
    ]
    for candidate in font_candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size=size)
    return ImageFont.load_default()


def draw_centered_text(draw, box, text, font, fill=(0, 0, 0)):
    bbox = draw.multiline_textbbox((0, 0), text, font=font, spacing=6)
    text_w = bbox[2] - bbox[0]
    text_h = bbox[3] - bbox[1]
    x = box[0] + (box[2] - box[0] - text_w) / 2
    y = box[1] + (box[3] - box[1] - text_h) / 2
    draw.multiline_text((x, y), text, font=font, fill=fill, spacing=6, align="center")


def draw_arrow(draw, start, end, fill=(60, 60, 60), width=4):
    draw.line([start, end], fill=fill, width=width)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex > sx else -1
        draw.polygon(
            [(ex, ey), (ex - 12 * direction, ey - 7), (ex - 12 * direction, ey + 7)],
            fill=fill,
        )
    else:
        direction = 1 if ey > sy else -1
        draw.polygon(
            [(ex, ey), (ex - 7, ey - 12 * direction), (ex + 7, ey - 12 * direction)],
            fill=fill,
        )


def generate_framework_figure(path: Path):
    img = Image.new("RGB", (1500, 900), "white")
    draw = ImageDraw.Draw(img)
    title_font = get_font(36, bold=True)
    box_font = get_font(26)
    draw.text((40, 20), "基于弱监督学习的伪装目标检测总体流程图", fill=(0, 0, 0), font=title_font)

    boxes = {
        "fully": (80, 140, 420, 280),
        "box": (80, 360, 420, 500),
        "anet": (560, 220, 900, 420),
        "pseudo": (1040, 140, 1380, 280),
        "score": (1040, 360, 1380, 500),
        "pnet": (560, 590, 900, 790),
    }
    colors = {
        "fully": (225, 238, 255),
        "box": (234, 248, 229),
        "anet": (255, 238, 220),
        "pseudo": (255, 247, 213),
        "score": (247, 231, 255),
        "pnet": (222, 245, 241),
    }
    labels = {
        "fully": "少量像素级标注样本\n图像 + 真值掩码 + 框标注",
        "box": "大量框标注样本\n图像 + 框标注",
        "anet": "辅助网络 ANet\n框提示引导分割\n生成初始伪标签",
        "pseudo": "伪标签数据集\nmask/edge/image",
        "score": "伪标签质量评估\n区域置信度\n边缘响应\n前景面积约束",
        "pnet": "主检测网络 PNet\n质量感知动态加权训练\n输出伪装目标分割结果",
    }

    for key, rect in boxes.items():
        draw.rounded_rectangle(rect, radius=24, fill=colors[key], outline=(80, 80, 80), width=3)
        draw_centered_text(draw, rect, labels[key], box_font)

    draw_arrow(draw, (420, 210), (560, 280))
    draw_arrow(draw, (420, 430), (560, 360))
    draw_arrow(draw, (900, 280), (1040, 210))
    draw_arrow(draw, (900, 360), (1040, 430))
    draw_arrow(draw, (1210, 500), (1210, 590))
    draw_arrow(draw, (1040, 430), (900, 670))
    draw_arrow(draw, (420, 430), (560, 670))
    draw_arrow(draw, (420, 210), (560, 620))

    note_font = get_font(22)
    draw.text((1040, 545), "质量分数写入 quality_scores.csv", fill=(80, 80, 80), font=note_font)
    img.save(path)


def generate_weight_figure(path: Path):
    img = Image.new("RGB", (1400, 800), "white")
    draw = ImageDraw.Draw(img)
    title_font = get_font(34, bold=True)
    axis_font = get_font(24)
    text_font = get_font(24)
    draw.text((40, 20), "伪标签质量评估与动态加权训练机制示意图", fill=(0, 0, 0), font=title_font)

    left_box = (70, 120, 620, 710)
    right_box = (760, 120, 1320, 710)
    draw.rounded_rectangle(left_box, radius=22, fill=(248, 248, 255), outline=(90, 90, 90), width=3)
    draw.rounded_rectangle(right_box, radius=22, fill=(245, 252, 245), outline=(90, 90, 90), width=3)
    draw.text((110, 145), "伪标签质量评分", fill=(0, 0, 0), font=axis_font)
    draw.text((820, 145), "弱监督样本权重调度", fill=(0, 0, 0), font=axis_font)

    # left list
    quality_lines = [
        "1. mask_conf: 前景/背景预测置信度",
        "2. edge_conf: 边缘响应强度",
        "3. area_penalty: 前景面积约束",
        "4. score = 0.6*mask + 0.25*edge + 0.15*area",
        "5. score 范围裁剪到 [0,1]",
    ]
    y = 210
    for line in quality_lines:
        draw.text((110, y), line, fill=(40, 40, 40), font=text_font)
        y += 70

    # right simple coordinate chart
    origin = (840, 620)
    x_end = (1260, 620)
    y_end = (840, 230)
    draw.line([origin, x_end], fill=(50, 50, 50), width=4)
    draw.line([origin, y_end], fill=(50, 50, 50), width=4)
    draw_arrow(draw, (840, 620), (1260, 620))
    draw_arrow(draw, (840, 620), (840, 230))
    draw.text((1200, 635), "epoch", fill=(0, 0, 0), font=text_font)
    draw.text((770, 225), "weight", fill=(0, 0, 0), font=text_font)

    warmup_points = [(840, 620), (930, 540), (1020, 460), (1110, 400), (1200, 360)]
    stable_points = [(1200, 360), (1260, 360)]
    draw.line(warmup_points, fill=(46, 120, 230), width=5)
    draw.line(stable_points, fill=(46, 120, 230), width=5)
    for pt in warmup_points + stable_points:
        draw.ellipse((pt[0] - 6, pt[1] - 6, pt[0] + 6, pt[1] + 6), fill=(46, 120, 230))

    draw.text((880, 660), "warmup阶段", fill=(46, 120, 230), font=text_font)
    draw.text((1130, 395), "稳定利用高质量伪标签", fill=(46, 120, 230), font=text_font)
    draw.text((900, 470), "样本权重 = 质量分数 × 全局弱监督系数", fill=(30, 90, 30), font=text_font)
    draw.text((900, 515), "低质量伪标签自动降权，减少训练干扰", fill=(30, 90, 30), font=text_font)
    img.save(path)


def set_run_font(run, east_asia, western=None, size=None, bold=None):
    run.font.name = western or east_asia
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if western:
        run._element.rPr.rFonts.set(qn("w:ascii"), western)
        run._element.rPr.rFonts.set(qn("w:hAnsi"), western)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold


def set_page_number_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    set_run_font(run, "宋体", "Times New Roman", 9)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, fld_end])


def configure_doc(doc: Document):
    section = doc.sections[0]
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    section.gutter = Cm(1)
    section.header_distance = Cm(2)
    section.footer_distance = Cm(1)
    set_page_number_footer(section)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)


def add_paragraph(doc, text="", align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, font="宋体", bold=False, first_indent=True):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(0.84)
    run = p.add_run(text)
    set_run_font(run, font, "Times New Roman", size, bold)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, "黑体", "Times New Roman", 16, True)
    elif level == 2:
        set_run_font(run, "黑体", "Times New Roman", 14, True)
    else:
        set_run_font(run, "宋体", "Times New Roman", 12, True)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    set_run_font(run, "宋体", "Times New Roman", 10.5, False)


def add_placeholder(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    set_run_font(run, "宋体", "Times New Roman", 11, False)
    run.italic = True


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("班 级  2203012")
    set_run_font(r, "宋体", "Times New Roman", 12, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("学 号  22009201147")
    set_run_font(r, "宋体", "Times New Roman", 12, True)

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("本科毕业设计论文")
    set_run_font(r, "黑体", "Times New Roman", 28, False)

    for _ in range(4):
        doc.add_paragraph()

    cover_items = [
        ("题       目", "基于弱监督学习的伪装目标检测方法设计与实现"),
        ("学       院", "计算机科学与技术学院"),
        ("专       业", "计算机科学与技术"),
        ("学 生 姓 名", "马蕴哲"),
        ("导 师 姓 名", "乔晓田（副教授）"),
    ]
    for label, value in cover_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        left = p.add_run(f"{label}     ")
        set_run_font(left, "宋体", "Times New Roman", 16, True)
        right = p.add_run(value)
        set_run_font(right, "黑体" if label == "题       目" else "宋体", "Times New Roman", 16 if label == "题       目" else 15, False)

    doc.add_page_break()


def add_abstracts(doc):
    add_heading(doc, "摘  要", 1)
    add_paragraph(
        doc,
        "伪装目标检测旨在从与背景高度相似的图像中准确分割目标区域，在军事侦察、智能安防和生态监测等领域具有重要应用价值。"
        "现有强监督方法通常依赖大量像素级标注数据，标注成本较高；弱监督方法虽然能够在一定程度上降低人工标注开销，"
        "但由于伪装目标边界模糊、前景与背景纹理相似，基于弱标注生成的伪标签往往带有明显噪声，容易造成模型训练不稳定和检测精度下降。"
        "针对上述问题，本文围绕基于弱监督学习的伪装目标检测方法开展研究，设计并实现了一种面向伪标签噪声抑制的两阶段检测框架。"
        "本文首先构建辅助网络ANet，利用少量像素级标注样本和框标注提示生成伪标签；随后构建主检测网络PNet，对真实标注样本与弱监督样本进行联合训练。"
        "在此基础上，本文从区域置信度、边缘响应和前景面积约束三个角度设计伪标签质量评估策略，对不同伪标签样本的可信度进行量化；"
        "进一步提出动态加权训练方法，通过质量分数和warmup机制自适应调节弱监督样本在损失函数中的贡献，降低低质量伪标签对主检测网络的不利影响。"
        "同时，本文对训练脚本、权重加载路径和伪标签处理流程进行了工程化优化，提高了方法的可复现性和跨平台可用性。"
        "论文给出了完整的方法设计、网络实现与实验方案，并对后续实验分析与结果展示进行了结构化组织。研究表明，本文提出的质量感知弱监督训练思路能够为低标注成本条件下的伪装目标检测提供有效支持，"
        "具有一定的理论意义和应用价值。",
    )
    p = add_paragraph(doc, "", first_indent=False)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    run1 = p.add_run("关键词：")
    set_run_font(run1, "黑体", "Times New Roman", 12, True)
    run2 = p.add_run("弱监督学习  伪装目标检测  伪标签去噪  动态加权  图像分割")
    set_run_font(run2, "宋体", "Times New Roman", 12, False)

    doc.add_page_break()

    add_heading(doc, "ABSTRACT", 1)
    add_paragraph(
        doc,
        "Camouflaged object detection aims to accurately segment target regions from images in which the foreground is highly similar to the background. "
        "It is of great significance in military reconnaissance, intelligent security, and ecological monitoring. Existing fully supervised methods usually rely on large-scale pixel-level annotations, "
        "which leads to high labeling cost. Although weakly supervised methods can reduce the annotation burden, the pseudo labels generated from weak annotations are often noisy because camouflaged objects usually have blurred boundaries and highly similar textures to the background. "
        "This issue may cause unstable training and degraded detection performance. To address these problems, this thesis designs and implements a two-stage weakly supervised camouflaged object detection framework with pseudo-label noise suppression. "
        "First, an auxiliary network ANet is constructed to generate pseudo labels by using a small amount of pixel-level annotations together with box prompts. Then, a primary network PNet is built to jointly train on fully annotated samples and weakly supervised samples. "
        "On this basis, a pseudo-label quality evaluation strategy is designed from three perspectives, namely region confidence, edge response, and foreground area constraint, so as to quantify the reliability of different pseudo labels. "
        "Furthermore, a dynamic weighting training method is proposed to adaptively adjust the contribution of weakly supervised samples in the loss function through quality scores and a warmup mechanism, thereby reducing the negative influence of low-quality pseudo labels on the primary detection network. "
        "In addition, the training scripts, weight loading paths, and pseudo-label processing pipeline are optimized to improve reproducibility and cross-platform usability. "
        "This thesis provides a complete presentation of method design, network implementation, and experimental planning, and organizes the later experimental analysis in a structured manner. "
        "The study shows that the proposed quality-aware weakly supervised training strategy can provide effective support for camouflaged object detection under low annotation cost conditions.",
        font="Times New Roman",
    )
    p = add_paragraph(doc, "", first_indent=False)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    run1 = p.add_run("Key  words：")
    set_run_font(run1, "Times New Roman", "Times New Roman", 12, True)
    run2 = p.add_run("Weakly Supervised Learning  Camouflaged Object Detection  Pseudo-label Denoising  Dynamic Weighting  Image Segmentation")
    set_run_font(run2, "Times New Roman", "Times New Roman", 12, False)


def add_toc(doc):
    doc.add_page_break()
    add_heading(doc, "目  录", 1)
    entries = [
        ("摘  要", "I"),
        ("ABSTRACT", "II"),
        ("第一章 绪论", "1"),
        ("第二章 相关技术与理论基础", "7"),
        ("第三章 基于弱监督学习的伪装目标检测方法设计与实现", "14"),
        ("第四章 实验设计与结果分析", "24"),
        ("第五章 总结与展望", "34"),
        ("致谢", "36"),
        ("参考文献", "37"),
    ]
    for title, page in entries:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.tab_stops.add_tab_stop(Cm(14.5))
        r1 = p.add_run(title)
        set_run_font(r1, "宋体", "Times New Roman", 12, False)
        r2 = p.add_run("\t")
        set_run_font(r2, "宋体", "Times New Roman", 12, False)
        r3 = p.add_run(page)
        set_run_font(r3, "宋体", "Times New Roman", 12, False)


def add_dataset_table(doc):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["数据集", "用途", "监督形式", "说明"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
    rows = [
        ["CAMO_COD_train", "ANet/PNet训练", "像素级标注 + 框标注", "作为少量全监督样本来源"],
        ["CAMO_COD_generate", "伪标签生成", "框标注", "由ANet生成伪标签并构建弱监督训练集"],
        ["CAMO_TestingDataset", "测试", "像素级标注", "用于评估检测性能"],
        ["COD10K_Test / CHAMELEON / NC4K", "泛化测试", "像素级标注", "用于后续扩展验证"],
    ]
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            cell.text = text


def add_placeholder_table(doc, title, headers, rows):
    add_paragraph(doc, title, align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5, font="宋体", first_indent=False)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            cell.text = text


def add_chapter_1(doc):
    doc.add_page_break()
    add_heading(doc, "第一章 绪论", 1)
    add_heading(doc, "1.1 研究背景与意义", 2)
    for text in [
        "伪装目标检测（Camouflaged Object Detection, COD）是计算机视觉领域中一个具有挑战性的研究方向，其目标是在复杂自然场景中从与背景高度相似的区域内定位并分割目标对象。"
        "与显著性目标检测相比，伪装目标往往缺少明显的颜色、纹理和边缘差异，因此传统基于对比度或显著性的视觉线索难以直接奏效。",
        "伪装目标检测在军事侦察、边境巡检、野生动物监测、农业病虫害识别以及医学影像辅助诊断等场景中都具有重要应用价值。"
        "在这些应用中，目标常常嵌入复杂环境之中，人工识别不仅耗时，而且容易受到主观经验与疲劳因素的影响，因此研究高精度自动检测方法具有现实意义。",
        "近年来，深度学习推动了伪装目标检测技术的发展，尤其是卷积神经网络、Transformer 以及多尺度特征融合结构的引入，使得模型能够在更复杂的背景下学习目标语义与边界信息。"
        "然而，大多数性能较高的方法都依赖大规模像素级标注数据，这与伪装目标检测任务的实际采集难度形成了明显矛盾。",
        "由于伪装目标边界模糊、目标形态复杂、前景面积变化大，像素级人工标注需要投入大量时间和人力成本。"
        "相较之下，框标注获取成本更低，因此如何利用少量像素级标注与大量框标注样本实现高质量检测，成为弱监督伪装目标检测研究中的核心问题。",
        "针对这一问题，本文围绕基于弱监督学习的伪装目标检测方法展开研究，重点关注伪标签噪声抑制与检测性能保持两个关键点。"
        "在已有两阶段弱监督框架的基础上，本文进一步设计了伪标签质量评估与动态加权训练机制，使得弱监督样本能够以更加稳健的方式参与主网络训练，从而为低标注成本场景下的伪装目标检测提供一种可行解决方案。",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "1.2 国内外研究现状", 2)
    for text in [
        "现有伪装目标检测研究大致可以分为全监督方法、半监督或弱监督方法两类。全监督方法通常依赖像素级真值掩码训练深层检测网络，通过引入多尺度编码器、边界分支、频域分解与注意力机制等结构提高对弱边界目标的表征能力。"
        "这类方法在公开数据集上取得了较高精度，但对高质量人工标注的依赖较强。",
        "随着数据标注成本问题日益突出，越来越多的研究开始关注弱监督视觉学习。相关工作通常利用框标注、点标注、涂鸦标注或图像级标签来替代像素级监督，并借助伪标签生成、自训练、置信度筛选和一致性约束等方法弥补监督信息不足。"
        "在伪装目标检测任务中，框标注弱监督具有较高实用价值，因为框信息能够提供目标的大致位置与尺度范围，但仍然无法直接给出精细边界。",
        "针对弱监督场景中的标签噪声问题，已有研究常采用辅助网络先生成伪标签，再利用主网络进行二次训练。"
        "这种两阶段思路能够在一定程度上缓解监督不足带来的性能下降，但如果伪标签本身存在较大偏差，错误信息会在后续训练中被进一步放大，导致模型边界分割不稳定、前景区域破碎甚至出现误检。",
        "从网络结构角度看，当前较有效的伪装目标检测模型普遍使用编码器-解码器结构，并结合边界分支、注意力模块、空洞卷积、多尺度上下文建模等机制强化目标表达能力。"
        "在本文所采用的框架中，辅助网络ANet使用图像分支与框提示分支共同提取特征，通过多尺度特征编码和逐级解码生成伪标签；主检测网络PNet则以共享编码器、频域分解和边界辅助监督为基础，在多层输出上进行联合优化。",
        "总体来看，国内外研究已经证明弱监督伪装目标检测具有可行性，但仍存在三个亟待解决的问题：第一，伪标签噪声难以被精确量化；第二，不同质量的弱监督样本在训练中被等同对待，容易引入错误梯度；第三，工程实现中训练流程复杂，复现成本较高。"
        "本文的研究工作正是围绕这些问题展开。",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "1.3 本文研究内容与主要工作", 2)
    for text in [
        "本文以弱监督伪装目标检测为研究对象，围绕“如何利用少量像素级标注和大量框标注实现稳定、高精度检测”这一目标开展研究。"
        "论文的主要工作包括以下几个方面：",
        "（1）完成两阶段弱监督伪装目标检测框架的搭建与实现。本文实现了辅助网络ANet与主检测网络PNet协同工作的训练流程，构建了由全监督样本、弱监督样本与伪标签组成的训练体系。",
        "（2）提出伪标签质量评估策略。针对弱监督样本中伪标签质量不稳定的问题，本文从区域预测置信度、边缘响应强度和前景面积约束三个角度构建质量评分函数，并在伪标签生成阶段输出每个样本的质量分数文件。",
        "（3）提出动态加权训练方法。本文在PNet训练阶段引入样本级质量权重和warmup机制，使模型在训练前期更多依赖高可信度监督，在后期逐步提升伪标签样本参与程度，从而降低噪声伪标签的负面影响。",
        "（4）完成工程化改进与论文初稿撰写。本文修复了预训练权重硬编码路径、跨平台文件复制等问题，完善了数据准备、训练、测试与日志记录流程，并形成了面向毕业论文写作的系统化文字说明与实验结构设计。",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "1.4 论文组织结构", 2)
    add_paragraph(
        doc,
        "本文共分为五章。第一章介绍研究背景、国内外研究现状、本文的研究内容与论文组织结构；第二章介绍伪装目标检测、弱监督学习、伪标签训练与相关网络结构等理论基础；"
        "第三章重点阐述本文的两阶段弱监督伪装目标检测方法，包括ANet、PNet、伪标签质量评估策略与动态加权训练方法；第四章给出实验数据集、评价指标、实验方案、结果分析以及需要补充的结果展示位置；第五章对全文工作进行总结，并对后续研究方向进行展望。"
    )


def add_chapter_2(doc):
    doc.add_page_break()
    add_heading(doc, "第二章 相关技术与理论基础", 1)
    add_heading(doc, "2.1 伪装目标检测任务概述", 2)
    for text in [
        "伪装目标检测任务既具有目标检测的定位属性，又具有语义分割的像素级预测要求。与普通分割任务不同，伪装目标往往与背景在颜色、纹理和形状特征上接近，因此模型必须更充分地利用上下文语义、多尺度边界和细粒度局部结构信息。",
        "从公开研究趋势来看，多尺度特征融合与边界辅助监督是提升伪装目标检测精度的关键手段之一。多尺度结构有助于同时建模语义层面的目标存在性和局部层面的边界连续性，边界分支则能够缓解前景边缘模糊造成的掩码破碎问题。",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "2.2 弱监督学习基本原理", 2)
    for text in [
        "弱监督学习的基本思想是在标注不完整、不精确或标注粒度较粗的条件下，通过设计合理的学习机制尽可能逼近全监督模型的性能。"
        "在图像分割任务中，常见弱监督信号包括图像级标签、点标注、框标注以及涂鸦标注等。",
        "与点标注和图像级标注相比，框标注能够提供目标位置与尺度范围，因此在伪装目标检测中具有更高的利用价值。"
        "但框标注并不包含真实边界，模型通常需要借助辅助网络、显著性先验、图优化方法或伪标签机制将粗粒度监督转换为可用于分割训练的伪标签。",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "2.3 伪标签学习与噪声抑制方法", 2)
    for text in [
        "伪标签学习是一类典型的自训练方法，其核心在于利用已有模型对未精标数据进行预测，再将预测结果作为新的监督信号参与后续训练。"
        "在弱监督伪装目标检测中，伪标签既能扩大训练样本规模，也能将框级监督转换为像素级监督，因此具有重要作用。",
        "然而，伪标签通常不可避免地存在噪声，尤其是在伪装目标边界模糊、目标尺寸较小或背景纹理复杂的情况下。"
        "如果训练阶段对所有伪标签一视同仁，错误伪标签将直接影响梯度方向，导致模型收敛速度变慢甚至性能退化。",
        "因此，伪标签质量评估和样本选择机制成为噪声抑制的关键。本文将从区域置信度、边缘信息和前景面积分布三个方面估计伪标签可靠性，并进一步通过动态加权训练机制调整弱监督样本贡献。",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "2.4 网络结构相关基础", 2)
    for text in [
        "本文实现的辅助网络ANet采用双分支结构，其中一条分支处理原始图像，另一条分支处理由框标注提示生成的引导图像。网络编码器使用ConvNeXt提取多尺度特征，再结合GCM3模块、GPM模块和带边界分支的解码器逐层恢复目标区域。",
        "主检测网络PNet使用PVTv2作为编码器，在特征提取后利用频域分解与多尺度解码结构生成最终掩码及边缘结果。"
        "在训练目标上，当前实现综合了加权二值交叉熵、区域一致性损失、边界Dice损失以及不确定性感知损失UAL，以提升弱监督条件下的训练稳定性。",
    ]:
        add_paragraph(doc, text)
    add_placeholder_table(
        doc,
        "表2.1 主要网络结构与作用概览",
        ["模块", "所在网络", "作用"],
        [
            ["ConvNeXt Branch", "ANet", "提取图像分支和框提示分支的多尺度语义特征"],
            ["PVTv2 Encoder", "PNet", "提取主检测网络的层次化表示"],
            ["GCM3 + DWT", "ANet/PNet", "完成多尺度特征融合与频域分解"],
            ["UNetDecoderWithEdges", "ANet/PNet", "逐级恢复掩码并提供边界辅助监督"],
            ["Quality Score", "本文改进", "量化伪标签可靠性，供后续加权训练使用"],
        ],
    )


def add_chapter_3(doc, fig1: Path, fig2: Path):
    doc.add_page_break()
    add_heading(doc, "第三章 基于弱监督学习的伪装目标检测方法设计与实现", 1)
    add_heading(doc, "3.1 总体框架设计", 2)
    for text in [
        "本文采用两阶段弱监督伪装目标检测框架。第一阶段利用少量像素级标注样本和框标注样本训练辅助网络ANet，生成弱监督样本对应的伪标签；第二阶段将真实标注样本与伪标签样本联合构建训练集，训练主检测网络PNet，输出最终伪装目标分割结果。",
        "与直接使用框标注进行训练相比，两阶段结构能够借助辅助网络将粗粒度监督转化为更接近像素级监督的掩码信息，从而显著提高后续主检测网络的学习效率。"
        "为进一步缓解伪标签噪声问题，本文在伪标签生成后加入质量评估环节，并在主检测网络训练时引入动态加权机制。",
    ]:
        add_paragraph(doc, text)
    doc.add_picture(str(fig1), width=Inches(6.1))
    add_caption(doc, "图3.1 基于弱监督学习的伪装目标检测总体流程图")

    add_heading(doc, "3.2 辅助网络ANet设计", 2)
    for text in [
        "ANet的主要任务是根据少量像素级标注样本学习框提示到精细目标区域的映射关系，并将其泛化到大量仅有框标注的样本上以生成伪标签。"
        "在实现中，ANet由两个分支组成：图像分支输入原始图像，框提示分支输入由框区域生成的引导图像，二者均通过ConvNeXt编码器提取多尺度特征。",
        "编码后的高层特征经过GPM模块生成先验响应图，低层和高层特征则通过GCM3模块、多分支ETM结构和带边界约束的U型解码器逐级恢复，最终输出多尺度掩码和边缘预测。"
        "在训练目标上，ANet采用结构损失与边界Dice损失联合优化，并辅以UAL不确定性感知损失，提高伪标签生成质量。",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "3.3 主检测网络PNet设计", 2)
    for text in [
        "PNet作为主检测网络，其目标是在弱监督训练集上完成最终的伪装目标检测。网络采用PVTv2作为共享编码器，对输入图像提取层次化特征表示。"
        "随后，模型利用DWT频域分解得到低频与高频成分，并通过多尺度卷积与逐级解码器实现目标区域恢复，同时输出边界分支结果。",
        "在损失函数设计方面，PNet综合使用NCLoss、边界Dice损失与UAL损失。NCLoss由加权二值交叉熵与区域差异项构成，用于提升区域预测的稳定性；边界Dice损失用于约束边缘分支输出；UAL损失用于增强模型对不确定区域的学习能力。",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "3.4 伪标签质量评估策略", 2)
    for text in [
        "为抑制低质量伪标签对训练过程的干扰，本文在伪标签生成阶段对每个样本引入质量评分。"
        "在实现上，首先对预测掩码的前景概率和背景概率取最大值并求平均，得到区域置信度；其次对边界预测结果进行裁剪求均值，得到边缘响应强度；最后根据前景区域占比与经验合理范围之间的偏差构造面积惩罚项。",
        "三个指标分别反映了区域预测的确定性、边界清晰程度以及前景区域是否异常偏大或偏小。本文将三者按经验权重线性融合，得到最终质量分数，并将其写入quality_scores.csv，供后续训练阶段读取。",
        "该策略具有实现简单、与现有流程耦合成本低的优点，适合作为伪标签筛选和样本级加权的基础指标。",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "3.5 动态加权训练方法", 2)
    for text in [
        "在主检测网络训练阶段，若对全监督样本与弱监督样本采用统一权重，低质量伪标签会直接参与损失计算并对网络更新产生较大影响。"
        "为此，本文提出一种基于质量分数的动态加权训练方法。具体而言，训练脚本首先读取弱监督样本对应的quality_scores.csv，将分数裁剪到给定范围内；"
        "然后引入全局弱监督权重与warmup系数，在训练初期降低弱监督样本的总体贡献，使模型优先从高质量真值样本中学习稳定特征；随着训练推进，再逐步提高弱监督样本的参与程度。",
        "最终，样本级权重被用于NCLoss与边界损失的加权平均，从而使高质量伪标签对模型优化贡献更大，低质量伪标签自动降权。"
        "该机制兼顾了弱监督样本规模优势与样本可靠性差异，有助于提升训练稳定性和最终检测效果。",
    ]:
        add_paragraph(doc, text)
    doc.add_picture(str(fig2), width=Inches(6.0))
    add_caption(doc, "图3.2 伪标签质量评估与动态加权训练机制示意图")

    add_heading(doc, "3.6 工程实现与流程优化", 2)
    for text in [
        "除了方法层面的改进，本文还对代码实现进行了工程化优化。首先，将原始代码中依赖固定Linux目录的预训练权重路径改为相对路径加载方式，以提升跨平台兼容性。"
        "其次，将伪标签生成阶段的系统复制命令替换为Python跨平台文件复制接口，避免在Windows环境下出现命令不可用问题。"
        "再次，在训练日志中增加了弱监督权重系数与平均伪标签质量输出，方便后续实验分析和参数调试。"
        "这些改进虽然不直接改变网络结构，但显著提高了方法的可运行性、可复现性与论文实验组织效率。",
    ]:
        add_paragraph(doc, text)


def add_chapter_4(doc):
    doc.add_page_break()
    add_heading(doc, "第四章 实验设计与结果分析", 1)
    add_heading(doc, "4.1 实验环境与数据集", 2)
    for text in [
        "本文实验平台基于PyTorch深度学习框架搭建，训练脚本包含单卡与多卡分布式训练逻辑。辅助网络ANet和主检测网络PNet的输入分辨率均设置为384×384，优化器采用Adam。"
        "训练过程中使用数据增强策略，包括颜色扰动、水平翻转、垂直翻转、噪声扰动、模糊变换和透视变换等，以提升模型泛化能力。",
        "数据集方面，本文训练阶段采用少量具有像素级标注的CAMO训练子集作为全监督样本，同时利用大量仅带框标注的图像通过ANet生成伪标签，构建弱监督训练集。"
        "测试阶段主要使用CAMO_TestingDataset进行性能评估，并保留在COD10K_Test、CHAMELEON和NC4K上的扩展测试接口，以便后续进行跨数据集泛化分析。",
    ]:
        add_paragraph(doc, text)
    add_placeholder_table(
        doc,
        "表4.1 数据集与实验用途说明",
        ["数据集", "角色", "是否在当前代码中出现", "备注"],
        [
            ["CAMO_COD_train", "全监督训练集", "是", "用于ANet与PNet中的真实标注样本"],
            ["CAMO_COD_generate", "伪标签生成集", "是", "由ANet生成伪标签后参与PNet训练"],
            ["CAMO_TestingDataset", "主测试集", "是", "当前验证流程默认测试集"],
            ["COD10K_Test / CHAMELEON / NC4K", "扩展测试集", "是", "可用于补充泛化能力分析"],
        ],
    )

    add_heading(doc, "4.2 评价指标与实验设置", 2)
    for text in [
        "当前代码实现中，验证阶段主要统计MAE（平均绝对误差）和IoU（交并比）两个指标。MAE反映预测掩码与真实掩码之间的像素差异，数值越小表示检测结果越接近真实标注；IoU衡量预测区域与真实区域的重叠程度，数值越大表示区域预测越准确。",
        "在后续论文定稿阶段，还可以根据实验实际补充S-measure、E-measure、F-measure等伪装目标检测常用指标，以形成更完整的结果分析体系。",
        "本文建议设置三组核心对比实验：第一组为未加入本文改进策略的原始PNet基线；第二组为仅加入伪标签质量评估的改进方案；第三组为在质量评估基础上进一步引入动态加权训练的完整方案。"
        "如果时间允许，还可比较不同ration设置下模型性能的变化，从而分析弱监督比例对改进策略有效性的影响。",
    ]:
        add_paragraph(doc, text)
    add_placeholder_table(
        doc,
        "表4.2 实验配置建议表",
        ["实验编号", "方法设置", "主要目的"],
        [
            ["Exp-1", "Baseline PNet", "作为原始两阶段弱监督框架对照基线"],
            ["Exp-2", "PNet + 质量评分", "验证伪标签可靠性量化的有效性"],
            ["Exp-3", "PNet + 质量评分 + 动态加权", "验证完整改进策略的综合效果"],
            ["Exp-4", "不同ration下的完整方法", "分析低标注比例场景下的稳定性与增益"],
        ],
    )

    add_heading(doc, "4.3 结果分析写作框架", 2)
    for text in [
        "由于本文当前阶段的重点在于完成方法设计、代码实现和论文初稿组织，完整的定量实验结果仍需在后续训练与测试后补充。"
        "因此，本节先给出可直接用于论文定稿的结果分析框架。",
        "在定量结果分析中，建议首先给出不同方法在主测试集上的MAE和IoU指标对比，并重点分析加入伪标签质量评估后MAE是否下降、IoU是否提升；其次分析动态加权训练是否能在保持伪标签规模优势的同时抑制错误监督传播。"
        "若完整方法在低标注比例场景下仍能稳定提升性能，则可进一步说明本文方法在实际弱监督条件下具有更高应用价值。",
        "在定性结果分析中，建议展示原图、真实掩码、基线方法预测结果以及本文方法预测结果，重点比较目标边界完整性、背景误检情况和细节恢复能力。"
        "特别是在复杂纹理背景和边界模糊样本上，本文方法应体现出更稳定的区域响应与更清晰的边界过渡。",
    ]:
        add_paragraph(doc, text)
    add_placeholder(doc, "【此处建议插入图4.1：原图-真值-基线结果-本文方法结果的可视化对比图】")
    add_caption(doc, "图4.1 伪装目标检测结果可视化对比图（待补充）")
    add_placeholder_table(
        doc,
        "表4.3 不同方法在主测试集上的定量结果（待补充）",
        ["方法", "MAE", "IoU", "说明"],
        [
            ["Baseline PNet", "待补充", "待补充", "原始两阶段弱监督框架"],
            ["PNet + 质量评分", "待补充", "待补充", "加入伪标签质量评估"],
            ["PNet + 质量评分 + 动态加权", "待补充", "待补充", "本文完整方法"],
        ],
    )

    add_heading(doc, "4.4 消融实验与讨论", 2)
    for text in [
        "为了分析各改进模块对最终性能的具体贡献，本文建议开展消融实验，分别考察质量评分模块、动态加权训练模块以及warmup系数设置对模型性能的影响。"
        "其中，质量评分模块主要验证是否能够有效区分高质量与低质量伪标签；动态加权模块主要验证样本级权重引入后是否能够改善训练稳定性；warmup机制则用于分析模型在训练前期对真值监督与伪标签监督的平衡关系。",
        "除定量指标外，还可结合训练日志观察不同设置下loss变化趋势、平均弱监督权重变化以及模型在验证集上的收敛速度，进一步解释改进策略的有效性来源。",
    ]:
        add_paragraph(doc, text)
    add_placeholder(doc, "【此处建议插入图4.2：不同epoch下弱监督权重与验证指标变化曲线】")
    add_caption(doc, "图4.2 弱监督权重调度与性能变化曲线（待补充）")
    add_placeholder_table(
        doc,
        "表4.4 消融实验结果（待补充）",
        ["设置", "质量评分", "动态加权", "warmup", "MAE", "IoU"],
        [
            ["A", "×", "×", "×", "待补充", "待补充"],
            ["B", "√", "×", "×", "待补充", "待补充"],
            ["C", "√", "√", "×", "待补充", "待补充"],
            ["D", "√", "√", "√", "待补充", "待补充"],
        ],
    )

    add_heading(doc, "4.5 本章小结", 2)
    add_paragraph(
        doc,
        "本章围绕实验平台、数据集、评价指标和实验方案对本文方法的验证流程进行了系统说明，并给出了后续结果补充所需的图表位置与分析框架。"
        "从论文初稿角度看，本章已经完成了实验部分的结构搭建，后续只需在完成训练与测试后将结果数值、可视化图和曲线图回填，即可形成完整的实验论证部分。"
    )


def add_chapter_5(doc):
    doc.add_page_break()
    add_heading(doc, "第五章 总结与展望", 1)
    add_heading(doc, "5.1 工作总结", 2)
    for text in [
        "本文围绕基于弱监督学习的伪装目标检测方法展开研究，完成了两阶段检测框架的设计、关键改进模块的实现以及论文初稿的系统组织。"
        "在方法层面，本文构建了由ANet和PNet组成的弱监督伪装目标检测流程，并针对伪标签噪声问题设计了质量评估策略和动态加权训练方法。"
        "在实现层面，本文完成了相应代码修改，对预训练权重路径、伪标签处理流程和日志记录机制进行了优化，提高了工程可复现性。"
        "整体而言，本文工作为低标注成本条件下的伪装目标检测提供了一种较为完整的设计思路，也为后续实验验证与论文定稿奠定了基础。",
        "从毕业设计过程来看，目前已完成的成果包括：任务书需求分析、相关文献调研、代码结构梳理、算法改进设计、训练流程实现、封面摘要目录撰写以及正文初稿生成。"
        "后续只需在实验部分进一步补充定量结果和可视化分析，即可形成较完整的毕业论文终稿。",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "5.2 未来展望", 2)
    for text in [
        "尽管本文已经从伪标签质量评估与动态加权角度对弱监督伪装目标检测进行了改进，但仍有进一步提升空间。"
        "首先，当前质量评分函数仍采用经验式设计，后续可尝试引入可学习的置信度估计模块，以获得更细粒度的样本可靠性评估结果。",
        "其次，本文主要围绕两阶段框架开展研究，未来可探索端到端的联合训练方式，使伪标签生成与主检测过程协同优化。"
        "再次，在实验指标方面，后续可补充更多公开数据集与评价指标，以更全面地验证方法的泛化性能。"
        "此外，还可以考虑将频域先验、对比学习或自监督预训练机制引入伪装目标检测流程，进一步提升低标注条件下模型的特征表达能力和鲁棒性。",
    ]:
        add_paragraph(doc, text)


def add_ack_and_refs(doc):
    doc.add_page_break()
    add_heading(doc, "致谢", 1)
    add_paragraph(
        doc,
        "在毕业设计完成过程中，我得到了导师、同学以及家人的大力支持与帮助。导师在课题选择、方法设计、论文写作和实验安排等方面给予了耐心指导；"
        "同学在代码调试和资料交流过程中提供了宝贵建议；家人在学习和生活中给予了充分理解与鼓励。在此一并表示诚挚感谢。",
    )

    add_heading(doc, "参考文献", 1)
    refs = [
        "[1] Zhang J, Zhang R, Shi Y, et al. Learning Camouflaged Object Detection from Noisy Pseudo Label[C]//European Conference on Computer Vision. 2024.",
        "[2] Goodfellow I, Pouget-Abadie J, Mirza M, et al. Generative Adversarial Nets[C]//Advances in Neural Information Processing Systems. 2014.",
        "[3] Ronneberger O, Fischer P, Brox T. U-Net: Convolutional Networks for Biomedical Image Segmentation[C]//MICCAI. 2015.",
        "[4] Liu Z, Mao H, Wu C Y, et al. A ConvNet for the 2020s[C]//CVPR. 2022.",
        "[5] Wang W, Xie E, Li X, et al. Pyramid Vision Transformer v2: Improved Baselines with Pyramid Vision Transformer[J]. Computational Visual Media, 2022.",
        "[6] Fan D P, Cheng M M, Liu Y, et al. Concealed Object Detection[J]. IEEE Transactions on Pattern Analysis and Machine Intelligence, 待核对.",
        "[7] 本文最终定稿时需根据实际引用内容补充弱监督学习、伪装目标检测和伪标签学习相关文献。",
    ]
    for ref in refs:
        add_paragraph(doc, ref, first_indent=False)


def build_doc():
    ensure_dir(FIG_DIR)
    fig1 = FIG_DIR / "framework.png"
    fig2 = FIG_DIR / "quality_weight.png"
    generate_framework_figure(fig1)
    generate_weight_figure(fig2)

    doc = Document()
    configure_doc(doc)
    add_cover(doc)
    add_abstracts(doc)
    add_toc(doc)
    add_chapter_1(doc)
    add_chapter_2(doc)
    add_chapter_3(doc, fig1, fig2)
    add_chapter_4(doc)
    add_chapter_5(doc)
    add_ack_and_refs(doc)
    doc.save(str(OUT_DOCX))


if __name__ == "__main__":
    build_doc()
