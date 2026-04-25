from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUTPUT_PATH = r"C:\Users\resii\Desktop\Noisy-cod-main-onlycode\thesis_cover_abstract_toc_draft.docx"


def set_run_font(run, east_asia, western=None, size=None, bold=None):
    run.font.name = western or east_asia
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if western:
        run._element.rPr.rFonts.set(qn("w:ascii"), western)
        run._element.rPr.rFonts.set(qn("w:hAnsi"), western)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold


def set_page_number_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    set_run_font(run, "宋体", "Times New Roman", 9)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, fld_end])


def set_document_layout(doc):
    section = doc.sections[0]
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    section.gutter = Cm(1)
    section.header_distance = Cm(2)
    section.footer_distance = Cm(1)
    set_page_number_footer(section)
    return section


def add_blank_paragraph(doc, text=""):
    p = doc.add_paragraph()
    if text:
        p.add_run(text)
    return p


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("班 级  2203012")
    set_run_font(r, "宋体", "Times New Roman", 12, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("学 号  22009201147")
    set_run_font(r, "宋体", "Times New Roman", 12, True)

    for _ in range(4):
        add_blank_paragraph(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("本科毕业设计论文")
    set_run_font(r, "黑体", "Times New Roman", 28, False)

    for _ in range(4):
        add_blank_paragraph(doc)

    cover_items = [
        ("题       目", "基于弱监督学习的伪装目标检测方法设计与实现"),
        ("学       院", "计算机科学与技术学院"),
        ("专       业", "计算机科学与技术"),
        ("学 生 姓 名", "马蕴哲"),
        ("导 师 姓 名", "乔晓田（副教授）"),
    ]

    for label, value in cover_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        label_run = p.add_run(f"{label}     ")
        set_run_font(label_run, "宋体", "Times New Roman", 16, True)
        value_run = p.add_run(value)
        set_run_font(value_run, "黑体" if label == "题       目" else "宋体", "Times New Roman", 16 if label == "题       目" else 15, False)

    doc.add_page_break()


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    r = p.add_run(text)
    set_run_font(r, "黑体", "Times New Roman", 16, True)


def add_body_paragraph(doc, text, english=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.first_line_indent = Cm(0.84)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    if english:
        set_run_font(r, "Times New Roman", "Times New Roman", 12, False)
    else:
        set_run_font(r, "宋体", "Times New Roman", 12, False)


def add_keywords(doc, label, content, english=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    label_run = p.add_run(label)
    if english:
        set_run_font(label_run, "Times New Roman", "Times New Roman", 12, True)
        content_run = p.add_run(content)
        set_run_font(content_run, "Times New Roman", "Times New Roman", 12, False)
    else:
        set_run_font(label_run, "黑体", "Times New Roman", 12, True)
        content_run = p.add_run(content)
        set_run_font(content_run, "宋体", "Times New Roman", 12, False)


def add_toc_line(doc, title, page):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = p.add_run(title)
    set_run_font(r1, "宋体", "Times New Roman", 12, False)
    r2 = p.add_run("\t")
    set_run_font(r2, "宋体", "Times New Roman", 12, False)
    r3 = p.add_run(str(page))
    set_run_font(r3, "宋体", "Times New Roman", 12, False)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(14.5))


def add_toc(doc):
    doc.add_page_break()
    add_heading(doc, "目  录")

    toc_entries = [
        ("摘  要", "I"),
        ("ABSTRACT", "II"),
        ("第一章 绪论", "1"),
        ("1.1 研究背景与意义", "1"),
        ("1.2 国内外研究现状", "2"),
        ("1.3 本文研究内容与主要工作", "5"),
        ("1.4 论文组织结构", "6"),
        ("第二章 相关技术与理论基础", "7"),
        ("2.1 伪装目标检测任务概述", "7"),
        ("2.2 弱监督学习基本原理", "9"),
        ("2.3 伪标签学习与噪声抑制方法", "11"),
        ("2.4 本章小结", "14"),
        ("第三章 基于弱监督学习的伪装目标检测方法设计", "15"),
        ("3.1 总体框架设计", "15"),
        ("3.2 辅助网络ANet设计", "17"),
        ("3.3 主检测网络PNet设计", "20"),
        ("3.4 伪标签质量评估策略", "23"),
        ("3.5 动态加权训练方法", "25"),
        ("3.6 本章小结", "27"),
        ("第四章 实验结果与分析", "28"),
        ("4.1 实验环境与数据集", "28"),
        ("4.2 评价指标与实验设置", "30"),
        ("4.3 定量实验结果分析", "32"),
        ("4.4 消融实验与可视化分析", "35"),
        ("4.5 本章小结", "38"),
        ("第五章 总结与展望", "39"),
        ("5.1 工作总结", "39"),
        ("5.2 未来展望", "40"),
        ("致谢", "41"),
        ("参考文献", "42"),
    ]

    for title, page in toc_entries:
        add_toc_line(doc, title, page)


def build_doc():
    doc = Document()
    set_document_layout(doc)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal_style.font.size = Pt(12)

    add_cover(doc)

    add_heading(doc, "摘  要")
    zh_abstract = (
        "伪装目标检测旨在从与背景高度相似的图像中准确分割目标区域，在军事侦察、"
        "智能安防和生态监测等领域具有重要应用价值。现有强监督方法通常依赖大量像素级标注数据，"
        "标注成本高；而弱监督场景下生成的伪标签往往存在较强噪声，容易导致模型训练不稳定与检测精度下降。"
        "针对上述问题，本文围绕基于弱监督学习的伪装目标检测方法展开研究，完成了方法设计、"
        "系统实现与实验流程优化。本文构建了由辅助网络ANet和主检测网络PNet构成的两阶段弱监督伪装目标检测流程，"
        "利用少量像素级标注和大量框标注样本生成伪标签，并构建弱监督训练数据集。在此基础上，设计了伪标签质量评估策略，"
        "从区域置信度、边缘响应和前景面积约束等角度对伪标签可靠性进行度量；进一步提出了面向弱监督样本的动态加权训练方法，"
        "通过质量分数与warmup机制降低低质量伪标签对主网络训练的干扰。同时，对预训练权重加载、伪标签复制与实验日志记录等工程细节进行了改进，"
        "提高了方法的可复现性与跨平台可用性。实验结果表明，本文设计的改进方法能够在弱监督场景下有效缓解伪标签噪声带来的负面影响，"
        "提升模型训练稳定性，并改善伪装目标检测效果。本文研究为低标注成本条件下的伪装目标检测提供了可行方案，具有一定的理论意义与应用价值。"
    )
    add_body_paragraph(doc, zh_abstract, english=False)
    add_keywords(doc, "关键词：", "弱监督学习  伪装目标检测  伪标签去噪  动态加权  图像分割", english=False)

    doc.add_page_break()

    add_heading(doc, "ABSTRACT")
    en_abstract = (
        "Camouflaged object detection aims to accurately segment target regions from images in which the targets are highly similar to the background, "
        "and it has important application value in military reconnaissance, intelligent security, and ecological monitoring. Existing fully supervised methods "
        "usually rely on a large amount of pixel-level annotations, resulting in high labeling cost. In weakly supervised settings, however, the generated pseudo labels "
        "are often noisy, which easily leads to unstable training and degraded detection accuracy. To address these issues, this thesis studies a weakly supervised camouflaged "
        "object detection method and completes the method design, system implementation, and experimental workflow optimization. "
        "First, a two-stage weakly supervised detection framework composed of ANet and PNet is constructed. A small amount of pixel-level annotations and a large amount of box annotations "
        "are used to generate pseudo labels and construct the weakly supervised training set. On this basis, a pseudo-label quality evaluation strategy is designed to measure the reliability "
        "of pseudo labels from the perspectives of region confidence, edge response, and foreground area constraint. Furthermore, a dynamic weighting training method for weakly supervised samples "
        "is proposed to reduce the interference of low-quality pseudo labels on the training of the main detection network through quality scores and a warmup mechanism. Meanwhile, engineering details "
        "such as pretrained weight loading, pseudo-label copying, and experiment logging are improved to enhance reproducibility and cross-platform usability. Experimental results show that the improved "
        "method can effectively alleviate the negative influence of noisy pseudo labels, improve training stability, and enhance camouflaged object detection performance under weak supervision. The study provides "
        "a feasible solution for camouflaged object detection with low annotation cost and has certain theoretical significance and practical value."
    )
    add_body_paragraph(doc, en_abstract, english=True)
    add_keywords(doc, "Key  words：", "Weakly Supervised Learning  Camouflaged Object Detection  Pseudo-label Denoising  Dynamic Weighting  Image Segmentation", english=True)

    add_toc(doc)

    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_doc()
